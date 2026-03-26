from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from api.errors import ApiError
from api.upload_parsing import build_upload_request, parse_upload_bytes
from packages.schemas.common.enums import InputComposition
from packages.schemas.output.error import ErrorCode


def make_docx_bytes(
    paragraphs: list[str],
    *,
    header: str | None = None,
    footer: str | None = None,
) -> bytes:
    document = Document()
    section = document.sections[0]
    if header is not None:
        section.header.paragraphs[0].text = header
    if footer is not None:
        section.footer.paragraphs[0].text = footer
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parse_upload_bytes_reads_utf8_text() -> None:
    content = parse_upload_bytes(filename="chapter.txt", content="第一章\n正文内容".encode("utf-8"))

    assert content == "第一章\n正文内容"


def test_parse_upload_bytes_extracts_only_docx_body_paragraphs() -> None:
    content = parse_upload_bytes(
        filename="chapter.docx",
        content=make_docx_bytes(
            ["第一段正文", "", "第二段正文"],
            header="页眉文本",
            footer="页脚文本",
        ),
    )

    assert content == "第一段正文\n\n第二段正文"
    assert "页眉文本" not in content
    assert "页脚文本" not in content


def test_build_upload_request_maps_chapters_file_to_single_chapter() -> None:
    request = build_upload_request(
        title="上传稿件",
        source_type="file_upload",
        chapters_text="第一章\n内容\n\n第二章\n内容",
        outline_text=None,
    )

    assert request.inputComposition is InputComposition.CHAPTERS_ONLY
    assert request.chapters is not None
    assert len(request.chapters) == 1
    assert request.chapters[0].content == "第一章\n内容\n\n第二章\n内容"


def test_parse_upload_bytes_rejects_oversized_content() -> None:
    with pytest.raises(ApiError) as exc_info:
        parse_upload_bytes(filename="chapter.txt", content=b"12345", max_bytes=4)

    assert exc_info.value.code is ErrorCode.UPLOAD_TOO_LARGE


def test_parse_upload_bytes_returns_parse_failed_for_invalid_docx() -> None:
    with pytest.raises(ApiError) as exc_info:
        parse_upload_bytes(filename="broken.docx", content=b"not-a-valid-docx")

    assert exc_info.value.code is ErrorCode.UPLOAD_PARSE_FAILED
